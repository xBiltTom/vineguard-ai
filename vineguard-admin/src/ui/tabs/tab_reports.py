"""
Pestaña de Reportes (PDF, Word, Excel).
"""

import os
import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
from fpdf import FPDF
import docx
from docx.shared import Pt, Inches

from src.locales.i18n import t
from src.core.statistics_core import create_beautiful_validation_charts, generate_interpretation_for_professor

def render():
    st.header("📄 Generación de Reportes Finales")
    
    # 1. Validar si existen datos dinámicos en memoria
    if 'mcnemar_analysis' not in st.session_state or st.session_state.mcnemar_analysis is None:
        st.warning("⚠️ No hay resultados estadísticos en memoria.")
        st.info("👉 Para generar un reporte real y dinámico, primero debes ir a la pestaña **3. Inferencia y Estadísticas**, cargar imágenes de prueba y ejecutar el Test de McNemar.")
        return

    st.markdown("""
    En esta sección puede descargar los reportes consolidados del entrenamiento, métricas de evaluación 
    y análisis estadístico (McNemar y MCC) basados **exactamente en los modelos y resultados que acabas de validar**.
    """)
    
    validation_data = st.session_state.mcnemar_validation
    mcnemar_analysis = st.session_state.mcnemar_analysis
    
    # Extraer datos reales
    y_true = validation_data['y_true']
    model_predictions = validation_data['predictions']
    model_names = validation_data['model_names']
    matthews_coefficients = mcnemar_analysis['matthews_coefficients']
    
    # Encontrar mejor modelo
    best_mcc_info = max(matthews_coefficients, key=lambda x: x['mcc'])
    best_model_name = best_mcc_info['model']
    
    # Construir DataFrame dinámico
    df_rows = []
    for i, (m_name, preds) in enumerate(zip(model_names, model_predictions)):
        acc = np.mean(y_true == preds) * 100
        mcc = matthews_coefficients[i]['mcc']
        df_rows.append({
            "Modelo": m_name,
            "Precisión (%)": round(acc, 2),
            "MCC": round(mcc, 3),
            "Recomendado": "Sí (Mejor MCC)" if m_name == best_model_name else "No"
        })
    
    df_metrics = pd.DataFrame(df_rows)
    
    # Generar gráficos reales y guardarlos temporalmente
    with st.spinner("Generando gráficos de alta calidad para los reportes..."):
        fig = create_beautiful_validation_charts(validation_data, mcnemar_analysis)
        chart_path = "temp_validation_chart.png"
        fig.savefig(chart_path, dpi=200, bbox_inches='tight')
        
        # Obtener interpretación experta
        interpretation = generate_interpretation_for_professor(mcnemar_analysis, validation_data)

    col1, col2, col3 = st.columns(3)
    
    # ==========================================
    # REPORTE EXCEL
    # ==========================================
    with col1:
        st.markdown("""
        <div class="tech-box" style="text-align: center;">
        <h2>📊</h2>
        <h4>Reporte Excel</h4>
        <p>Tablas completas de entrenamiento y métricas generadas dinámicamente.</p>
        </div>
        """, unsafe_allow_html=True)
        
        output_xl = BytesIO()
        with pd.ExcelWriter(output_xl, engine='xlsxwriter') as writer:
            df_metrics.to_excel(writer, sheet_name='Resultados_Métricas', index=False)
            workbook = writer.book
            worksheet = writer.sheets['Resultados_Métricas']
            header_format = workbook.add_format({'bold': True, 'bg_color': '#415D48', 'font_color': 'white'})
            for col_num, value in enumerate(df_metrics.columns.values):
                worksheet.write(0, col_num, value, header_format)
                worksheet.set_column(col_num, col_num, 20)
        excel_data = output_xl.getvalue()
        
        st.download_button("📥 Descargar Excel", data=excel_data, file_name="Reporte_VineGuard_Metricas.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
        
    # ==========================================
    # REPORTE WORD
    # ==========================================
    with col2:
        st.markdown("""
        <div class="tech-box" style="text-align: center;">
        <h2>📝</h2>
        <h4>Reporte Word</h4>
        <p>Documento con gráficas integradas e interpretación experta del MCC.</p>
        </div>
        """, unsafe_allow_html=True)
        
        doc = docx.Document()
        doc.add_heading('Reporte de Desempeño y Validación Estadística - VineGuard AI', 0)
        
        doc.add_heading('1. Gráficas de Validación Cruzada y Matrices', level=1)
        doc.add_picture(chart_path, width=Inches(6.5))
        
        doc.add_heading('2. Tabla de Métricas de Modelos', level=1)
        table = doc.add_table(rows=1, cols=len(df_metrics.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df_metrics.columns):
            hdr_cells[i].text = str(column_name)
        for index, row in df_metrics.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
                
        doc.add_heading('3. Conclusiones y Dictamen (McNemar)', level=1)
        # Limpiar markdown de la interpretación
        clean_interp = interpretation.replace('**', '')
        doc.add_paragraph(clean_interp)
        
        output_doc = BytesIO()
        doc.save(output_doc)
        word_data = output_doc.getvalue()
        
        st.download_button("📥 Descargar Word", data=word_data, file_name="Reporte_VineGuard_Interpretacion.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        
    # ==========================================
    # REPORTE PDF
    # ==========================================
    with col3:
        st.markdown("""
        <div class="tech-box" style="text-align: center;">
        <h2>📄</h2>
        <h4>Reporte PDF</h4>
        <p>Versión ejecutiva oficial con gráficos de alta resolución para la defensa.</p>
        </div>
        """, unsafe_allow_html=True)
        
        class PDF(FPDF):
            def header(self):
                self.set_font('Arial', 'B', 15)
                self.cell(0, 10, 'Reporte Ejecutivo Oficial - VineGuard AI', 0, 1, 'C')
                self.ln(5)
                
        pdf = PDF()
        pdf.add_page()
        
        # Tabla de métricas
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, txt="1. Resumen de Modelos Evaluados", ln=True)
        pdf.set_font("Arial", '', 10)
        for index, row in df_metrics.iterrows():
            texto = f"- {row['Modelo']} | Precisión: {row['Precisión (%)']}% | MCC: {row['MCC']} | Rec: {row['Recomendado']}"
            pdf.cell(0, 8, txt=texto, ln=True)
            
        pdf.ln(5)
        
        # Insertar imagen de gráficas
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, txt="2. Visualización Estadística (MCC, Matrices, Distribución)", ln=True)
        # Ancho 190 para abarcar casi toda la página A4
        pdf.image(chart_path, x=10, w=190)
        
        # Conclusiones en una nueva página si no cabe
        pdf.add_page()
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, txt="3. Dictamen y Conclusiones", ln=True)
        pdf.set_font("Arial", '', 10)
        # Limpiar caracteres UTF-8 no soportados por fpdf nativo
        clean_interp = interpretation.replace('**', '').replace('•', '-').replace('≥', '>=')
        pdf.multi_cell(0, 7, txt=clean_interp)
        
        pdf_data = pdf.output(dest='S').encode('latin1')
        
        st.download_button("📥 Descargar PDF", data=pdf_data, file_name="Reporte_VineGuard_Ejecutivo.pdf", mime="application/pdf", use_container_width=True)
        
    st.markdown("---")
    st.subheader("Vista Previa de Resultados Detectados Dinámicamente")
    st.dataframe(df_metrics, use_container_width=True)
    
    # Limpieza del archivo temporal
    if os.path.exists(chart_path):
        try:
            os.remove(chart_path)
        except:
            pass
